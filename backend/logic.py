import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import pyodbc
import logging

# 配置日誌
# 設定日誌級別為 INFO，並定義日誌格式，以便追蹤程式執行情況
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# 常量定義
# 定義表格標題的背景填充顏色 (灰色)
HEADER_FILL_COLOR = "D9D9D9"
# 定義表格內容行中，標題列的背景填充顏色 (淺灰色)
ROW_HEADER_FILL_COLOR = "F2F2F2"
# 當資料為 NaN 或空值時，預設顯示的字串
DEFAULT_NAN_DISPLAY = "NaN"
# 當 SQL 語法設定鍵值查無對應 SQL 語法時，預設顯示的字串
DEFAULT_SQL_NOT_FOUND = " 查無對應 SQL"
# 批次流程 ID 的前綴，用於篩選相關資料 (例如 'FI_%' 表示以 'FI_' 開頭的流程)
API_FLOW_PREFIX = 'FI_%'
# SQL Server 連線選項，用於信任伺服器憑證 (避免憑證錯誤)
TRUST_SERVER_CERTIFICATE = "yes" 

def load_sql_properties(path: Path) -> dict:
    """
    從指定的檔案路徑載入 SQL 查詢的屬性。
    檔案格式應為 key=value，並忽略以 '#' 開頭的註解行。
    """
    sql_dict = {}
    try:
        # 以讀取模式開啟檔案，使用 UTF-8 編碼
        with open(path, 'r', encoding='utf-8') as f:
            for line in f:
                # 檢查行中是否包含 '=' 且不是註解行 (以 '#' 開頭)
                if '=' in line and not line.strip().startswith('#'):
                    # 以第一個 '=' 分割鍵和值
                    key, val = line.split('=', 1)
                    # 移除鍵和值前後的空白字元後存入字典
                    sql_dict[key.strip()] = val.strip()
    except FileNotFoundError:
        # 如果檔案未找到，記錄錯誤並重新拋出
        logger.error(f"SQL 屬性檔案未找到: {path}")
        raise
    except Exception as e:
        # 捕獲其他載入時的異常，記錄錯誤並重新拋出
        logger.error(f"載入 SQL 屬性檔案時發生錯誤: {e}")
        raise
    return sql_dict

def generate_api_doc(sql_connection_params: dict, word_template_path: Path, output_path: Path, sql_properties_path: Path):
    """
    根據資料庫中的API資訊和SQL屬性檔案，生成API說明文件。

    Args:
        sql_connection_params (dict): 資料庫連線參數，包含 'server', 'database', 'username', 'password'。
        word_template_path (Path): Word 文件範本的路徑。
        output_path (Path): 輸出 Word 文件的路徑。
        sql_properties_path (Path): 包含 SQL 查詢語法的屬性檔案路徑。
    """
    # 從傳入的字典中獲取資料庫連線參數
    server = sql_connection_params['server']
    database = sql_connection_params['database']
    username = sql_connection_params['username']
    password = sql_connection_params['password']

    conn = None # 初始化資料庫連線物件為 None
    try:
        # 構建 ODBC 連線字串
        conn_str = (
            f'DRIVER={{ODBC Driver 18 for SQL Server}};' # 指定 ODBC 驅動程式
            f'SERVER={server};'                         # 伺服器地址
            f'DATABASE={database};'                     # 資料庫名稱
            f'UID={username};'                          # 使用者名稱
            f'PWD={password};'                          # 密碼
            f"TrustServerCertificate={TRUST_SERVER_CERTIFICATE};" # 信任伺服器憑證
        )
        logger.info(f"嘗試連接資料庫: Server={server}, Database={database}")
        conn = pyodbc.connect(conn_str) # 建立資料庫連線
        logger.info("資料庫連線成功。")

        # 獲取所有相關的 API 代碼，用於後續查詢的過濾條件
        # 這有助於避免在每個查詢中重複龐大的子查詢，提高效率
        relevant_api_codes_query = f"""
            SELECT DISTINCT Y.CALL_CODE_ID
            FROM JH_WS02_FLOW_LIST X
                INNER JOIN JH_WS02_FLOW_SCHEDULE_LIST Y ON X.pk = Y.FLOW_ID_PK
                INNER JOIN JH_WS02_CODE_LIST Z ON Y.CALL_CODE_ID = Z.CODE_ID
            WHERE FLOW_ID LIKE '{API_FLOW_PREFIX}'
        """
        relevant_api_codes_df = pd.read_sql(relevant_api_codes_query, conn)
        # 提取不重複的 API 代碼，並轉換為列表
        relevant_api_codes = relevant_api_codes_df['CALL_CODE_ID'].dropna().unique().tolist()
        
        # 如果沒有找到相關的 API 代碼，則直接生成一個簡單的 Word 文件並返回
        if not relevant_api_codes:
            logger.warning(f"資料庫中找不到 FLOW_ID 像 '{API_FLOW_PREFIX}' 的相關 API 代碼。")
            doc = Document(word_template_path)
            doc.add_paragraph(f"資料庫中找不到與 '{API_FLOW_PREFIX}' 相關的 API 資料。")
            doc.save(output_path)
            return

        # 將相關 API 代碼轉換為 SQL IN 子句所需的字串格式 (例如: "'CODE1', 'CODE2'")
        # 由於 pyodbc.read_sql 不直接支援 params 傳遞列表給 IN 子句，這裡直接構建字串
        # 注意：雖然此處 relevant_api_codes 來自資料庫而非用戶輸入，風險可控，但仍需注意 SQL 注入風險
        api_codes_str = ", ".join(f"'{code}'" for code in relevant_api_codes)

        # 查詢 API 階層表資訊 (批次代碼、說明、API 順序、代碼、說明)
        api_hierarchy_df = pd.read_sql(f"""
            SELECT
                a.FLOW_ID AS "批次代碼"
                ,a.FLOW_HELP AS "批次說明"
                ,class_num AS "API順序"
                ,b.CALL_CODE_ID AS "API代碼"
                ,c.API_DESC AS "API說明"
            FROM JH_WS02_FLOW_LIST a
                INNER JOIN JH_WS02_FLOW_SCHEDULE_LIST b ON a.pk = b.FLOW_ID_PK
                INNER JOIN JH_WS02_CODE_LIST c ON b.CALL_CODE_ID = c.CODE_ID
            WHERE FLOW_ID LIKE '{API_FLOW_PREFIX}'
            ORDER BY a.FLOW_HELP, CLASS_NUM
        """, conn)

        # 查詢所有 API 清單資訊 (代碼、簡述、說明、類型、連線名稱、執行類型、語法鍵值、驗證金鑰、是否編碼)
        api_list_df_all = pd.read_sql(f"""
            SELECT
                CODE_ID AS [API代碼]
                ,API_DESC AS [API簡述]
                ,REPLACE(REPLACE(CODE_HELP,CHAR(13)+CHAR(10),''),CHAR(10),'') AS [API說明]
                ,CASE EXEC_TYPE WHEN '0' THEN 'SQL' ELSE 'SSH' END AS [API行為類型]
                ,JNDI_USE AS [資料庫連線名稱]
                ,REPLACE(ACTION_TYPE,CHAR(13)+CHAR(10),'') AS [執行類型]
                ,REPLACE(SQL_PROP_KEY,CHAR(13)+CHAR(10),'') AS [語法設定鍵值]
                ,'DFMDB_authority' AS [驗證金鑰]
                ,CASE IS_ENCODE WHEN 'Y' THEN '是' ELSE '否' END AS [是否編碼]
            FROM JH_WS02_CODE_LIST
            WHERE CODE_ID IN ({api_codes_str})
        """, conn)

        # 查詢所有輸出設定資訊 (API 代碼、節點階層、父子鍵值、輸出參數)
        output_setting_df_all = pd.read_sql(f"""
            SELECT A.CODE_ID AS [API代碼]
                ,ISNULL(B.CLASS_NUM,'') AS [節點階層]
                ,ISNULL(B.UP_PK_FIELD,'') AS [父階層關聯鍵值]
                ,ISNULL(B.DOWN_PK_FIELD,'') AS [子階層關聯鍵值]
                ,ISNULL(B.OUTPUT_FIELD,'') AS [輸出參數]
            FROM JH_WS02_CODE_LIST A
                LEFT JOIN JH_WS02_CODE_FORMAT_LIST B ON A.PK = B.CODE_ID_PK
            WHERE CODE_ID IN ({api_codes_str})
            ORDER BY 1,2
        """, conn)

        # 查詢所有 IP 權限設定資訊 (API 代碼、IP、說明)
        ip_permission_df_all = pd.read_sql(f"""
            SELECT A.CODE_ID AS [API代碼]
                ,B.ACCESSED_IP AS [IP]
                ,ISNULL(B.ACCESSED_DESC,'') AS [說明]
            FROM JH_WS02_CODE_LIST A
                LEFT JOIN JH_WS02_CODE_IP_RELATION B ON A.PK = B.CODE_ID_PK
            WHERE CODE_ID IN ({api_codes_str})
            ORDER BY 1,2
        """, conn)

        # 查詢所有 WebService 資訊 (API 代碼、序、主機代碼、名稱、IP、啟用狀態)
        webservice_df_all = pd.read_sql(f"""
            SELECT A.CODE_ID AS [API代碼]
                ,B.CLASS_NUM AS [序]
                ,ISNULL(B.WEB_SERVICE_CODE,'') AS [主機代碼]
                ,CASE B.WEB_SERVICE_CODE WHEN 'WS01' THEN 'Middle01' WHEN 'WS02' THEN 'Middle02' END AS [主機名稱]
                ,CASE B.WEB_SERVICE_CODE WHEN 'WS01' THEN '192.168.222.136' WHEN 'WS02' THEN '192.168.222.138' END AS [主機IP]
                ,CASE B.IS_DOING WHEN 'Y' THEN '是' ELSE '否' END AS [啟用]
            FROM JH_WS02_CODE_LIST A
                LEFT JOIN JH_WS02_CODE_WS_RELATION B ON A.PK = B.CODE_ID_PK
            WHERE CODE_ID IN ({api_codes_str})
            ORDER BY 1,2
        """, conn)

        # 查詢所有參數驗證資訊 (API 代碼、序、屬性名、預設值、說明)
        param_validation_df_all = pd.read_sql(f"""
            SELECT A.CODE_ID AS [API代碼]
                ,ROW_NUMBER() OVER (PARTITION BY A.CODE_ID ORDER BY B.FORMAT_IDX) AS [序]
                ,ISNULL(B.INPUT_FIELD,'') AS [屬性名]
                ,ISNULL(B.INPUT_DEFAULT_VAL,'') AS [預設值]
                ,ISNULL(REPLACE(B.REG_DESC,CHAR(13)+CHAR(10),'') ,'') AS [說明]
            FROM JH_WS02_CODE_LIST A
                LEFT JOIN JH_WS02_CODE_RANGE_ANALYSIS B ON A.PK = B.CODE_ID_PK
            WHERE CODE_ID IN ({api_codes_str})
            ORDER BY 1,2
        """, conn)

        api_data = {} # 用於儲存每個 API 代碼相關的所有資料 (字典，鍵為 API 代碼，值為該 API 的各類別資料 DataFrame)

        # 輔助函數，用於將 DataFrame 按 'API代碼' 填充到 api_data 字典
        def populate_api_data(df: pd.DataFrame, sheet_name: str):
            for api_code in df['API代碼'].dropna().unique():
                if api_code not in api_data:
                    api_data[api_code] = {} # 如果 API 代碼不存在，則創建一個新的字典
                api_data[api_code][sheet_name] = df[df['API代碼'] == api_code] # 儲存該 API 代碼對應的資料
                
        # 呼叫輔助函數填充 api_data 字典
        populate_api_data(api_list_df_all, 'API清單')
        populate_api_data(output_setting_df_all, '輸出設定')
        populate_api_data(ip_permission_df_all, 'IP權限設定')
        populate_api_data(webservice_df_all, 'WebService')
        populate_api_data(param_validation_df_all, '參數驗證')

        # 載入 SQL 屬性檔案（包含實際的 SQL 語法）
        sql_map = load_sql_properties(sql_properties_path)
        logger.info(f"成功載入 SQL 屬性檔案: {sql_properties_path}")

        # 載入 Word 範本文件
        doc = Document(word_template_path)
        logger.info(f"成功載入 Word 範本: {word_template_path}")

        # 遍歷批次清單，為每個批次生成文件內容
        # 提取不重複的批次代碼和說明，並按批次說明排序
        batch_list = api_hierarchy_df[['批次代碼', '批次說明']].drop_duplicates().sort_values(by=['批次說明'])
        for row in batch_list.itertuples(index=False): # 迭代每一行，不包含索引
            batch_code, batch_desc = row # 解構批次代碼和說明
            
            # 根據批次代碼和說明篩選相關的 API 階層資料
            group_df = api_hierarchy_df[
                (api_hierarchy_df['批次代碼'] == batch_code) &
                (api_hierarchy_df['批次說明'] == batch_desc)
            ].copy() # 使用 .copy() 避免 SettingWithCopyWarning (修改副本而不是原始 DataFrame 的視圖)

            # 在 Word 文件中添加批次標題
            doc.add_paragraph(f'{batch_code} ({batch_desc})', style='Heading 2')
            
            # API 階層表 (固定顯示，即使沒有 API 也會顯示表頭)
            api_list_table = doc.add_table(rows=1, cols=3) # 創建一個 1 行 3 列的表格
            api_list_table.style = 'Table Grid' # 設定表格樣式為網格
            api_list_table.autofit = False # 關閉自動調整寬度
            headers = ['順序', 'API代碼', 'API說明'] # 表頭名稱
            widths_cm = [1.24, 7, 10.79] # 每列的寬度 (公分)

            # 設定表頭內容、寬度及背景顏色
            for i in range(3):
                cell = api_list_table.cell(0, i) # 獲取表頭單元格
                cell.text = headers[i] # 設定單元格文字
                cell.width = Cm(widths_cm[i]) # 設定單元格寬度
                tcPr = cell._tc.get_or_add_tcPr() # 獲取或添加單元格屬性
                shd = OxmlElement('w:shd') # 創建背景填充元素
                shd.set(qn('w:fill'), HEADER_FILL_COLOR) # 設定填充顏色
                tcPr.append(shd) # 將填充元素添加到單元格屬性

            # 處理 API 階層表資料，檢查每個值是否為 NaN 或空字串
            if group_df.empty: # 如果篩選後的資料框為空
                cells = api_list_table.add_row().cells # 添加一行並獲取所有單元格
                cells[0].text = DEFAULT_NAN_DISPLAY
                cells[1].text = DEFAULT_NAN_DISPLAY
                cells[2].text = DEFAULT_NAN_DISPLAY
                for i in range(3):
                    cells[i].width = Cm(widths_cm[i]) # 保持寬度
                logger.info(f"批次 '{batch_code}' ('{batch_desc}') 下無 API 階層資料。")
            else:
                for _, api_row in group_df.iterrows(): # 遍歷篩選後的資料框中的每一行
                    cells = api_list_table.add_row().cells # 添加一行並獲取所有單元格
                    # 獲取各欄位值，若不存在則預設為空字串
                    seq_val = api_row.get('API順序', '')
                    api_code_val = api_row.get('API代碼', '')
                    api_desc_val = api_row.get('API說明', '')

                    # 格式化並設定單元格文字，處理 NaN 和空字串
                    cells[0].text = str(int(seq_val)) if pd.notna(seq_val) and float(seq_val).is_integer() else DEFAULT_NAN_DISPLAY
                    cells[1].text = str(api_code_val) if pd.notna(api_code_val) and str(api_code_val).strip() != '' else DEFAULT_NAN_DISPLAY
                    cells[2].text = str(api_desc_val) if pd.notna(api_desc_val) and str(api_desc_val).strip() != '' else DEFAULT_NAN_DISPLAY
                    for i in range(3):
                        cells[i].width = Cm(widths_cm[i]) # 保持寬度

            doc.add_paragraph('') # 添加一個空段落作為間隔

            if group_df.empty:
                continue # 如果此批次沒有 API 資料，則跳過後續的詳細處理

            # 遍歷每個批次下的 API，生成詳細資訊
            for _, api_row in group_df.iterrows():
                api_code = api_row['API代碼']
                # 檢查 API 代碼是否為空，若為空則跳過
                if pd.isna(api_code) or str(api_code).strip() == '':
                    logger.warning(f"批次 '{batch_code}' 下發現空的 API 代碼，跳過處理。")
                    continue
                
                doc.add_paragraph(api_code, style='Heading 4') # 添加 API 代碼作為小標題
                
                # API 清單表格的表頭設定
                api_detail_table = doc.add_table(rows=1, cols=3)
                api_detail_table.style = 'Table Grid'
                api_detail_table.autofit = False
                headers_api_list = ['序', '參數', '設定值']
                widths_cm_api_list = [1.24, 3.5, 14.29]

                # 設定 API 清單表格的表頭
                for i in range(3):
                    cell = api_detail_table.cell(0, i)
                    cell.text = headers_api_list[i]
                    cell.width = Cm(widths_cm_api_list[i])
                    tcPr = cell._tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), HEADER_FILL_COLOR)
                    tcPr.append(shd)

                # 從 api_data 字典中獲取當前 API 代碼對應的 'API清單' 資料
                api_df = api_data.get(api_code, {}).get('API清單')
                
                # 如果 API 清單資料存在且不為空
                if api_df is not None and not api_df.empty:
                    # 定義 API 清單中要顯示的參數名稱及其順序
                    param_names = [
                        'API代碼', 'API簡述', 'API說明', 'API行為類型', '資料庫連線名稱',
                        '執行類型', '語法設定鍵值', '驗證金鑰', '是否編碼', '語法'
                    ]
                    for i, param in enumerate(param_names, start=1): # 遍歷每個參數名稱
                        row_cells = api_detail_table.add_row().cells # 添加一行並獲取所有單元格
                        row_cells[0].text = str(i) # 設定序號
                        row_cells[1].text = param # 設定參數名稱
                        
                        value_to_display = DEFAULT_NAN_DISPLAY # 預設顯示 NaN

                        if param == '語法': # 特殊處理 '語法' 欄位
                            config_key_val = api_df['語法設定鍵值'].iloc[0] # 獲取語法設定鍵值
                            if pd.notna(config_key_val) and str(config_key_val).strip() != '':
                                # 從 sql_map 中獲取對應的 SQL 語法，如果沒有則顯示查無對應 SQL
                                raw_value = sql_map.get(str(config_key_val), DEFAULT_SQL_NOT_FOUND)
                                # 替換換行符號和 tab 符號，處理 escape 字符
                                value_to_display = raw_value.replace('\\n', '\n').replace('\\t', '\t').replace('\\=', '=')
                            logger.debug(f"API {api_code} - 語法鍵值: {config_key_val}, 語法: {value_to_display[:50]}...")
                        else: # 其他一般參數
                            if param in api_df.columns: # 如果參數存在於資料框的欄位中
                                val_from_df = api_df[param].iloc[0] # 獲取對應的值
                                if pd.notna(val_from_df) and str(val_from_df).strip() != '':
                                    value_to_display = str(val_from_df) # 如果值不為 NaN 且不為空，則使用該值

                        row_cells[2].text = value_to_display # 設定設定值
                        for j in range(3):
                            row_cells[j].width = Cm(widths_cm_api_list[j]) # 保持寬度
                        # 設定前兩列的背景顏色
                        for j in [0, 1]:
                            tcPr = row_cells[j]._tc.get_or_add_tcPr()
                            shd = OxmlElement('w:shd')
                            shd.set(qn('w:fill'), ROW_HEADER_FILL_COLOR)
                            tcPr.append(shd)
                else: # 如果 API 清單資料為空
                    row_cells = api_detail_table.add_row().cells
                    row_cells[0].text = DEFAULT_NAN_DISPLAY
                    row_cells[1].text = DEFAULT_NAN_DISPLAY
                    row_cells[2].text = DEFAULT_NAN_DISPLAY
                    for j in range(3):
                        row_cells[j].width = Cm(widths_cm_api_list[j])
                    logger.warning(f"API {api_code} 缺少 'API清單' 資料。")

                doc.add_paragraph('') # 添加空段落

                # 定義其他 API 相關區塊的順序、表頭和寬度
                section_order = [
                    ('參數驗證', ['序', '屬性名', '預設值', '說明'], [1.24, 3.5, 8.89, 5.4]),
                    ('WebService', ['序', '主機代碼', '主機名稱', '主機IP', '啟用'], [1.24, 3.5, 4.57, 4.32, 5.4]),
                    ('IP權限設定', ['IP', '說明'], [8.74, 10.25]),
                    ('輸出設定', ['節點階層', '父階層關聯鍵值', '子階層關聯鍵值', '輸出參數'], [3.24, 4.5, 5.89, 5.4])
                ]
                
                # 遍歷每個區塊並生成其表格
                for sheet_name, headers, widths_cm in section_order:
                    doc.add_paragraph(sheet_name, style='Heading 5') # 添加區塊標題
                    table = doc.add_table(rows=1, cols=len(headers)) # 創建表格
                    table.style = 'Table Grid'
                    table.autofit = False
                    
                    # 設定表格的表頭
                    for i, header in enumerate(headers):
                        cell = table.cell(0, i)
                        cell.text = header
                        cell.width = Cm(widths_cm[i])
                        tcPr = cell._tc.get_or_add_tcPr()
                        shd = OxmlElement('w:shd')
                        shd.set(qn('w:fill'), HEADER_FILL_COLOR)
                        tcPr.append(shd)

                    # 從 api_data 字典中獲取當前 API 代碼對應的區塊資料
                    df = api_data.get(api_code, {}).get(sheet_name)
                    
                    # 如果區塊資料存在且不為空
                    if df is not None and not df.empty:
                        for _, row_data in df.iterrows(): # 遍歷每一行資料
                            row_cells = table.add_row().cells # 添加一行
                            for j, header in enumerate(headers): # 遍歷每一列
                                value = row_data.get(header, '') # 從 DataFrame 取得值，預設為空字串

                                display_value = DEFAULT_NAN_DISPLAY # 預設顯示 "NaN"

                                # 如果值不是 NaN 且去除前後空白後不為空字串
                                if pd.notna(value) and str(value).strip() != '':
                                    # 特殊處理 '序' 和 '節點階層' 欄位，轉換為整數顯示
                                    if header in ['序', '節點階層']:
                                        try:
                                            # 如果是浮點數且為整數值，則轉換為整數
                                            display_value = str(int(value)) if float(value).is_integer() else str(value)
                                        except ValueError:
                                            display_value = str(value) # 無法轉換則保持原字串
                                    else:
                                        display_value = str(value) # 其他欄位直接使用字串形式

                                row_cells[j].text = display_value # 設定單元格文字
                                row_cells[j].width = Cm(widths_cm[j]) # 設定單元格寬度
                    else:
                        # 如果此區塊的 DataFrame 為空，在表格中添加一行並填入 "NaN"
                        row_cells = table.add_row().cells
                        for j in range(len(headers)):
                            row_cells[j].text = DEFAULT_NAN_DISPLAY
                            row_cells[j].width = Cm(widths_cm[j])
                        logger.info(f"API {api_code} 缺少 '{sheet_name}' 資料。")
                    doc.add_paragraph('') # 添加空段落

        # 儲存最終的 Word 文件
        doc.save(output_path)
        logger.info(f"API 規格書成功生成於: {output_path}")

    except pyodbc.Error as e:
        # 捕獲 pyodbc 資料庫錯誤，記錄並重新拋出
        logger.error(f"資料庫操作失敗: {e}")
        raise # 重新拋出異常，以便 main.py 捕獲並返回給前端
    except FileNotFoundError as e:
        # 捕獲檔案未找到錯誤，記錄並重新拋出
        logger.error(f"檔案未找到錯誤: {e}")
        raise
    except Exception as e:
        # 捕獲其他所有未預期的錯誤，記錄詳細的堆棧追蹤
        logger.exception(f"文件生成過程中發生未知錯誤: {e}") # 使用 exception 會記錄堆棧追蹤
        raise
    finally:
        # 確保在函式結束時關閉資料庫連線
        if conn:
            conn.close()
            logger.info("資料庫連線已關閉。")